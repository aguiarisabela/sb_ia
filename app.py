from flask import Flask, render_template, request, send_file
from transformers import T5Tokenizer, T5ForConditionalGeneration
from docx import Document
import os

app = Flask(__name__)

tokenizer = T5Tokenizer.from_pretrained("t5-base")
model = T5ForConditionalGeneration.from_pretrained("t5-base")

def traduzir_texto(texto):
    input_text = f"translate English to Portuguese: {texto}"
    
   
    inputs = tokenizer(input_text, return_tensors="pt", padding=True, truncation=True)
    outputs = model.generate(inputs["input_ids"], max_length=512, num_beams=4, early_stopping=True)

    
    traducao = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return traducao


def traduzir_docx(file_path):
    doc = Document(file_path)
    translated_doc = Document()

    
    for para in doc.paragraphs:
        traduzido = traduzir_texto(para.text)
        translated_doc.add_paragraph(traduzido)
    
    
    translated_file_path = "translated_file.docx"
    translated_doc.save(translated_file_path)
    return translated_file_path

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/traduzir_texto', methods=['POST'])
def traduzir_texto_web():
    texto = request.form['texto']
    texto_traduzido = traduzir_texto(texto)
    return render_template('index.html', texto_traduzido=texto_traduzido)

@app.route('/traduzir_docx', methods=['POST'])
def traduzir_docx_web():
    file = request.files['file']
    file_path = os.path.join("uploads", file.filename)
    file.save(file_path)
    

    translated_file_path = traduzir_docx(file_path)
    

    return send_file(translated_file_path, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
